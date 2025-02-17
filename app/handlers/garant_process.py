import requests
import random
from datetime import datetime
from bs4 import BeautifulSoup
from docx import Document
import mysql.connector
from mysql.connector import Error
import os
import logging


# Константы
TOKEN = "9d97b394982611ef8cf20050568d72f0"
DOCUMENTS_FOLDER_HTML = "documents_html"
DOCUMENTS_FOLDER_DOCX = "documents_docx"
MYSQL_CONFIG = {
    "host": "194.87.243.188",
    "port": 3306,
    "user": "gen_user",
    "password": "63)$0oJ\WRP\$J",
    "database": "default_db"
}


def process_garant_request(query, logs, rag_module):
    """Обрабатывает запрос к API Гаранта."""
    try:
        headers = {"Authorization": f"Bearer {TOKEN}"}
        payload = {
            "text": query,
            "count": 1,
            "kind": ["003"],  # Судебная практика
            "sort": 0,
            "sortOrder": 0
        }
        response = requests.post("https://api.garant.ru/v1/search", headers=headers, json=payload)
        response.raise_for_status()
        data = response.json()

        if not data.get("documents"):
            logs.append(rag_module("error", "Документы не найдены."))
            return None

        document = data["documents"][0]
        topic_id = document["topic"]
        document_name = document["name"]
        logs.append(rag_module("info", f"Документ найден: {document_name} (topic_id: {topic_id})."))

        html_content = download_html(topic_id, logs, rag_module)
        if not html_content:
            logs.append(rag_module("error", "Не удалось скачать HTML контент документа."))
            return None

        docx_file_path = convert_html_to_docx(html_content, topic_id, logs, rag_module)
        document_url = f"http://127.0.0.1:5000/download/{topic_id}.docx"

        save_document_to_db(
            document_id=str(random.randint(10000000, 99999999)),
            document_name=document_name.split("Ключевые темы")[0].strip(),
            document_num=topic_id,
            document_url=document_url,
            download_date=datetime.now(),
            logs=logs,
            rag_module=rag_module
        )

        return {"docx_file_path": docx_file_path, "document_url": document_url}

    except requests.exceptions.RequestException as e:
        logs.append(rag_module("error", f"Ошибка при запросе к API Гаранта: {e}"))
        return None


def download_html(topic_id, logs, rag_module):
    """Скачивает HTML контент для указанного topic_id через API Гаранта."""
    html_file_path = f"{DOCUMENTS_FOLDER_HTML}/{topic_id}.html"
    os.makedirs(DOCUMENTS_FOLDER_HTML, exist_ok=True)

    try:
        headers = {"Authorization": f"Bearer {TOKEN}"}
        response = requests.get(f"https://api.garant.ru/v1/topic/{topic_id}/html", headers=headers, verify=False)
        response.raise_for_status()
        json_data = response.json()

        if "items" not in json_data or not json_data["items"]:
            logs.append(rag_module("error", "Ответ не содержит валидного HTML."))
            return None

        html_content = json_data["items"][0]["text"]
        with open(html_file_path, "w", encoding="utf-8") as file:
            file.write(html_content)

        logs.append(rag_module("info", f"HTML успешно сохранён: {html_file_path}"))
        return html_content

    except requests.exceptions.RequestException as e:
        logs.append(rag_module("error", f"Ошибка при скачивании HTML: {e}"))
        return None


def convert_html_to_docx(html_content, topic_id, logs, rag_module):
    """Конвертирует HTML-контент в DOCX."""
    docx_file_path = f"{DOCUMENTS_FOLDER_DOCX}/{topic_id}.docx"
    os.makedirs(DOCUMENTS_FOLDER_DOCX, exist_ok=True)

    try:
        document = Document()
        soup = BeautifulSoup(html_content, "html.parser")

        for element in soup.find_all(["p", "h1", "h2", "h3", "h4", "h5", "h6", "strong", "em"]):
            text = element.get_text(strip=True)
            if element.name.startswith("h"):
                level = int(element.name[1])
                document.add_heading(text, level=level)
            elif element.name == "strong":
                run = document.add_paragraph().add_run(text)
                run.bold = True
            elif element.name == "em":
                run = document.add_paragraph().add_run(text)
                run.italic = True
            else:
                document.add_paragraph(text)

        document.save(docx_file_path)
        logs.append(rag_module("info", f"DOCX файл сохранен: {docx_file_path}."))
        return docx_file_path

    except Exception as e:
        logs.append(rag_module("error", f"Ошибка при конвертации HTML в DOCX: {e}"))
        return None



def save_document_to_db(document_id, document_name, document_num, document_url, download_date, logs, rag_module):
    """Сохраняет информацию о документе в базу данных."""
    conn = None
    cursor = None  # Добавляем явную инициализацию, чтобы избежать проблем

    try:
        conn = mysql.connector.connect(**MYSQL_CONFIG)

        if conn.is_connected():
            cursor = conn.cursor()

            query = """
                INSERT INTO documents (document_id, document_name, document_num, document_url, download_date)
                VALUES (%s, %s, %s, %s, %s)
            """
            cursor.execute(query, (document_id, document_name, document_num, document_url, download_date))
            conn.commit()

            logs.append(rag_module("info", f"✅ Информация о документе сохранена в БД: {document_id}."))

    except Error as e:
        error_msg = f"❌ Ошибка базы данных: {e}"
        logs.append(rag_module("error", error_msg))
        logging.error(error_msg)

    finally:
        if cursor:
            cursor.close()  # Закрываем курсор перед закрытием соединения
        if conn and conn.is_connected():
            conn.close()
            logging.info("✅ Подключение к базе данных закрыто.")

